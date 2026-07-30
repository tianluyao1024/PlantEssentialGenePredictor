"""Create a The Plant Cell-oriented independent-validation manuscript draft.

The draft reports only the actual result of the frozen external-evidence gate:
the current curated cohort is below the pre-registered size/class minimum, so
it deliberately contains no external AUC/AUPRC or threshold-performance claim.
The source document remains unchanged.
"""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


SOURCE = Path(r"C:\Users\tly\Downloads\Plant_essential_gene_prediction_manuscript_web_deployment_Nature_style.docx")
OUTPUT = Path(r"C:\Users\tly\Downloads\Plant_essential_gene_prediction_manuscript_ThePlantCell_independent_validation_draft.docx")
FIGURE = Path(r"E:\PlantEssentialGenePredictor\results\manuscript_figures\figure6_audited_candidate_resource\Figure6_audited_candidate_resource.png")
FIGURE7 = Path(r"E:\PlantEssentialGenePredictor\results\manuscript_figures\figure7_external_evidence_boundary\Figure7_external_evidence_boundary.png")


def set_text(paragraph, text: str) -> None:
    """Replace a paragraph while retaining its paragraph style."""
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)


def insert_after(paragraph, text: str = "", style: str | None = None):
    new_p = paragraph._parent.add_paragraph(style=style)
    paragraph._p.addnext(new_p._p)
    if text:
        new_p.add_run(text)
    return new_p


def insert_before(paragraph, text: str = "", style: str | None = None):
    new_p = paragraph._parent.add_paragraph(style=style)
    paragraph._p.addprevious(new_p._p)
    if text:
        new_p.add_run(text)
    return new_p


def find_paragraph(document: Document, prefix: str):
    for paragraph in document.paragraphs:
        if paragraph.text.startswith(prefix):
            return paragraph
    raise ValueError(f"paragraph not found: {prefix}")


def add_heading_after(paragraph, text: str):
    p = insert_after(paragraph, style="Heading 2")
    p.add_run(text)
    return p


def add_body_after(paragraph, text: str):
    p = insert_after(paragraph, style="Normal")
    p.add_run(text)
    return p


def add_caption_after(paragraph, text: str):
    p = insert_after(paragraph, style="Caption")
    p.add_run(text)
    return p


def add_figure_after(paragraph, figure: Path):
    p = insert_after(paragraph, style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(figure), width=Inches(6.9))
    return p


def main() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    for figure in (FIGURE, FIGURE7):
        if not figure.exists():
            raise FileNotFoundError(figure)
    doc = Document(SOURCE)

    set_text(
        doc.paragraphs[0],
        "A leakage-aware cross-species framework prioritizes plant essential-gene candidates in Arabidopsis and rice",
    )
    set_text(
        doc.paragraphs[2],
        "The Plant Cell-oriented manuscript draft with frozen model evaluation, an audited candidate resource and a bounded independent-evidence analysis.",
    )
    abstract = find_paragraph(doc, "Essential genes are central")
    set_text(
        abstract,
        "Essential genes support plant viability, development and fertility, yet genome-scale experimental essentiality maps remain incomplete in plants. Here we developed a leakage-aware framework that integrates conservative phenotype-label curation with a common 6,751-dimensional representation of Arabidopsis thaliana and Oryza sativa genes. The representation combines 95 biological features with ESM2, ProtBERT and ProtT5 protein-language-model embeddings. On fixed held-out test sets, species-specific models achieved area under the receiver-operating-characteristic curve (AUC) values of 0.9175 in Arabidopsis and 0.8812 in rice. A joint model gave numerically higher rice AUC but no significant paired DeLong improvement, and modestly lower Arabidopsis AUC. Feature ablation and homology-cluster evaluation showed complementary biological and sequence-representation signal while exposing annotation dependence and sequence-similarity effects. To separate discovery from label reuse, we audited all feature-covered genes against study labels, pseudo-labels and raw phenotype-record sources. The resulting release contains 17,522 Arabidopsis and 17,457 rice true-unknown candidates, alongside a curated core set of ten candidates per species. A targeted external source screen generated 312 articles and 919 gene-level records; 19 curator-checked direct loss-of-function records passed the zero-overlap audit. The pre-registered cohort-size requirement was not met, so no external performance estimate is reported. We provide locked splits, model assets, code, candidate audit and a public prediction interface for evidence-based candidate prioritization rather than replacement of experimental confirmation.",
    )

    web_heading = find_paragraph(doc, "A web-accessible predictor supports")
    set_text(web_heading, "Public prediction interface and reusable candidate resource")
    web_body = find_paragraph(doc, "For raw-data workflows")
    set_text(
        web_body,
        "The public interface supports a reproducible raw-data workflow. Users upload protein and CDS FASTA files and may add GFF3, GO annotation, PPI edge-list, expression-matrix and domain-annotation files through distinct validated upload fields. The system preserves gene identifiers, reports feature coverage, retains the longest protein-coding transcript per gene and selects only a released model compatible with the available input profile. Raw uploads are temporary; only a user-authorized final species-level prediction table is eligible for public caching.",
    )
    web_body_2 = find_paragraph(doc, "To make the released models accessible")
    set_text(
        web_body_2,
        "PlantEssentialGenePredictor is implemented as a public Streamlit application (https://plantessentialgene.com). It provides species-specific Arabidopsis and rice full models, a joint Arabidopsis-rice model and deployable feature-profile models for incomplete annotations. The web application is a delivery mechanism for frozen models and auditable outputs; it is not used as biological validation of model predictions.",
    )

    interpretation_caption = find_paragraph(doc, "Figure 5. Feature-interpretation summary")
    h1 = add_heading_after(interpretation_caption, "Audited genome-scale candidate release")
    b1 = add_body_after(
        h1,
        "Genome-scale prediction tables were reclassified before candidate nomination. Each feature-covered gene was assigned exactly one release status: known_label_used_in_study, pseudo_label_used_in_study, phenotype_recorded_but_excluded or true_unknown_candidate. A candidate could enter the true-unknown class only if it had zero overlap with the strict modeling labels, Arabidopsis pseudo-labels where applicable, and every reconciled raw phenotype-record source. The automated exclusion audit confirmed zero prohibited overlaps for all released core candidates.",
    )
    b2 = add_body_after(
        b1,
        "This procedure retained 17,522 true-unknown Arabidopsis genes and 17,457 true-unknown rice genes. We then nominated ten core candidates per species using concordant essential calls at each frozen model's validation-derived threshold and a probability of at least 0.50 in all three frozen species-specific, joint and annotation-light model outputs. Core candidates were additionally ranked by ensemble-component robustness and screened against the closest labelled homolog. For each species, the displayed panel deliberately includes five candidates with and five without a stringent labelled-essential homolog. These strata are intended to distinguish candidates supported by local sequence neighbourhoods from candidates that are not readily explained by a close labelled homolog.",
    )
    figure_p = add_figure_after(b2, FIGURE)
    fig_caption = add_caption_after(
        figure_p,
        "Figure 6. Audited genome-scale candidate resource. (a) Feature-covered genes were separated into study labels, study pseudo-labels, genes with phenotype records that were excluded from the primary labels and true-unknown candidates. (b,c) Agreement of frozen species-specific, joint and annotation-light prediction families for the ten Arabidopsis and ten rice core candidates. The dashed line marks probability 0.5 and is shown only as a visual reference. (d) Closest-labelled-homolog audit. A candidate was called homology-supported only when its closest study-labelled homolog met the predefined identity and coverage thresholds. These panels describe nomination evidence, not independent phenotype validation.",
    )
    h2 = add_heading_after(fig_caption, "Prespecified independent phenotype-validation framework")
    b3 = add_body_after(
        h2,
        "Candidate-level functional annotations, GO terms, PPI attributes and expression summaries can help formulate mechanistic hypotheses, but they are not treated as independent validation when the same source family contributed to the feature matrix. We therefore screened public literature and open mutant records with a source ledger that records the query, date, stable identifier, normalized gene ID, inclusion or exclusion decision and source-level provenance. Each direct phenotype record was required to be absent from the frozen study-label registry, pseudo-label registry and every archived raw phenotype source before allele-level manual adjudication.",
    )
    b4 = add_body_after(
        b3,
        "The initial curated set comprised 19 direct Arabidopsis loss-of-function records, of which 16 met all phenotype-adjudication requirements for a prelocked cohort (nine essential and seven viable/non-essential). All 19 passed the automated zero-overlap audit. However, the preregistered requirement was at least 30 genes per species and at least ten genes per class. The cohort therefore failed the gate, and the frozen evaluator withheld external AUC, AUPRC, sensitivity, specificity, precision and F1. The same screening identified no curator-locked rice record. Thus, neither species is assigned an external performance estimate in this manuscript.",
    )
    figure7_p = add_figure_after(b4, FIGURE7)
    figure7_caption = add_caption_after(
        figure7_p,
        "Figure 7. Independent-evidence boundary and qualitative candidate cards. (a) Europe PMC literature screening, automated exclusion against the frozen study registry and manual allele-level direct loss-of-function curation. (b) The pre-registered external-cohort gate. The 16 curated Arabidopsis records (nine essential and seven viable/non-essential) and zero curator-locked rice records did not meet the required n >= 30 per species and n >= 10 per class; external discrimination and threshold metrics were therefore not calculated. (c) Two Arabidopsis core candidates with at least two independent evidence categories. DG409 (AT1G01970) has direct CRISPR embryo-lethality evidence and functional organellar evidence. MISF74 (AT4G01400) is retained as a qualitative counterexample because viable but severely growth-retarded insertion mutants were reported. (d) Evidence coverage of frozen core candidates. Prediction-only entries are not presented as biologically validated examples.",
    )
    h3 = add_heading_after(figure7_caption, "Independent evidence is informative but not yet a second performance test")
    add_body_after(
        h3,
        "The locked evidence set nevertheless illustrates the intended use of the resource. For DG409 (AT1G01970), a 2023 CRISPR study reported embryo lethality, and experimental work linked the encoded PPR protein to chloroplast and mitochondrial development. In contrast, MISF74 (AT4G01400), a high-scoring core candidate, has published viable insertion mutants with severe growth retardation and mitochondrial intron-splicing defects. We report this counterexample rather than selectively presenting only concordant cases. Eight Arabidopsis and all ten rice core candidates remain prediction-only resource entries because they do not yet satisfy the requirement for at least two independent evidence categories.",
    )

    discussion_first = find_paragraph(doc, "This study supports three main conclusions")
    set_text(
        discussion_first,
        "This study supports three main conclusions. First, plant essential-gene ranking is feasible when phenotype-derived labels are curated conservatively and evaluated on locked splits. Second, protein-language-model representations and interpretable biological features provide complementary signal, while ablations reveal important annotation dependence that must be reported rather than hidden. Third, genome-scale predictions must be separated from the labels and phenotype records used to build the model. The audited release therefore distinguishes true-unknown candidates from study labels, pseudo-labels and phenotype-recorded but excluded genes, and treats candidate nomination as hypothesis generation rather than biological confirmation.",
    )
    discussion_anchor = find_paragraph(doc, "A second limitation is annotation bias")
    d1 = add_body_after(
        discussion_anchor,
        "The candidate resource is organized around a biological working model in which low redundancy, conserved core molecular complexes and developmentally constrained programs jointly increase the probability of essentiality. The homology-supported and homology-isolated strata were retained to prevent this model from being reduced to simple nearest-neighbour annotation transfer. The present data do not justify a strong claim of zero-shot transfer across plant species: joint training altered the sensitivity-specificity trade-off but did not significantly improve AUC by paired testing.",
    )
    add_body_after(
        d1,
        "The next evidentiary step is expansion of the independently curated phenotype cohort or targeted mutant validation that is genuinely independent of all model-development decisions. The present source screen demonstrates both the value and the difficulty of this requirement: most literature-derived gene records were excluded because they overlapped historical phenotype archives, lacked stable gene-to-allele resolution or described compound/conditional genotypes. The released source ledger, automated independence checks and evidence-card schema make expansion reproducible. In the interim, public predictions should be used to prioritize experiments, combine orthogonal evidence and guide resource allocation rather than to replace phenotypic assays.",
    )

    method_anchor = find_paragraph(doc, "ROC AUC and AUPRC")
    mh1 = add_heading_after(method_anchor, "Candidate release audit and core-candidate nomination")
    mb1 = add_body_after(
        mh1,
        "For every feature-covered gene, identifiers were normalized to the study gene namespace and compared with strict modeling-label tables, the Arabidopsis 0.60/0.40 pseudo-label table and each raw phenotype-record source used during data collection. Genes were assigned a mutually exclusive publication status in the following order: known_label_used_in_study, pseudo_label_used_in_study, phenotype_recorded_but_excluded, or true_unknown_candidate. Only the final class was eligible for the candidate resource. An automated audit reports candidate intersections with every prohibited source; the required intersection count is zero.",
    )
    mb2 = add_body_after(
        mb1,
        "For each true-unknown gene, frozen species-specific, joint and annotation-light predictors were evaluated without refitting. Core candidates were selected only when all three models produced an essential call at their locked validation-derived thresholds and each probability was at least 0.50; they were then ranked by component-model robustness and stratified by DIAMOND similarity to the closest study-labelled gene. Homology-supported candidates met both the predefined identity and query-coverage criteria against a labelled essential gene; homology-isolated candidates did not. This screen is a sequence-neighbourhood audit and is not an external validation test.",
    )
    mh2 = add_heading_after(mb2, "Independent phenotype-validation protocol")
    add_body_after(
        mh2,
        "The external-cohort evaluator accepts only phenotype records that include a stable source identifier, publication or release date, experiment type, developmental stage, binary label under a prespecified rule and an explicit independence declaration. It rejects records overlapping model labels, pseudo-labels or phenotype-recorded-but-excluded genes, including all raw phenotype archives used during data collection. The cohort must be frozen before predictions are inspected. Quantitative evaluation is enabled only for a species with at least 30 genes and at least ten genes in each class; otherwise an underpowered status report is written and no metric is calculated. If enabled in a later release, continuous probabilities will be evaluated without retraining, thresholded metrics will use the already locked validation-derived threshold and 95% confidence intervals will be computed by 10,000 stratified bootstrap resamples.",
    )

    availability = find_paragraph(doc, "Code, processed label tables")
    set_text(
        availability,
        "Code, lightweight release tables, fixed train/validation/test splits, feature lists and figure-generation scripts are available at https://github.com/tianluyao1024/PlantEssentialGenePredictor, including the frozen independent-validation supplement at https://github.com/tianluyao1024/PlantEssentialGenePredictor/releases/tag/v1.2.3-independent-validation. Processed feature matrices, trained model assets, complete genome-scale prediction tables, source-screening ledgers, curated-record audits, prelocked cohort status reports, evidence-card tables, figure source data and frozen-input checksums are archived on Zenodo v1.2.3 (https://doi.org/10.5281/zenodo.21387076; concept DOI https://doi.org/10.5281/zenodo.21387075). The web interface is available at https://plantessentialgene.com.",
    )

    # Insert a concise supplementary-data list before the reference heading.
    reference_heading = find_paragraph(doc, "References")
    supplement_heading = insert_before(reference_heading, style="Heading 1")
    supplement_heading.add_run("Supplementary data release")
    s1 = add_body_after(
        supplement_heading,
        "Supplementary Data 1: audited Arabidopsis genome-scale prediction table with mutually exclusive publication status. Supplementary Data 2: audited rice genome-scale prediction table with mutually exclusive publication status. Supplementary Data 3: Arabidopsis and rice core-candidate evidence cards and evidence-card summaries. Supplementary Data 4: automated candidate-exclusion audit. Supplementary Data 5: Europe PMC source-screening ledger and curated external source-screening ledger. Supplementary Data 6: curator-checked direct loss-of-function records, zero-overlap audit and prelocked external-cohort status report. Supplementary Data 7: Figure 6 and Figure 7 source data, release manifest and reproducibility pipeline.",
    )
    s1.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.core_properties.title = "Leakage-aware plant essential-gene candidate prioritization"
    doc.core_properties.subject = "The Plant Cell-oriented manuscript revision"
    doc.core_properties.comments = "Independent source screening and the pre-registered underpowered external-cohort result added; no external performance metric is fabricated."
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
